import os
from dotenv import load_dotenv
import streamlit as st
from openai import AzureOpenAI
import time
# Azure OpenAI
import pandas as pd
from datetime import datetime, date
from docx import Document
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

from process_flow import show_process_flow
from send_email import send_jd_email


from io import BytesIO
from docx import Document
import streamlit as st

def create_docx(jd_text):
    doc = Document()
    doc.add_heading("Job Description", level=1)
    doc.add_paragraph(jd_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def save_request(
    next_req_id,
    requester_name,
    requester_empid,
    requested_for_name,
    requested_for_empid,
    req_job_title,
    req_job_code,
    onshore_offshore,
    expected_hire_date,
    direct_reports,
    posting_scope,
    hiring_need,
    business_justification,
    job_family,
    level,
    business_unit,
    quarter,
    doc_path
):
    request_row = {
        "Department": business_unit,
        "RequesterName": requester_name,
        "RequestedEmpID": requester_empid,
        "RequestedForName": requested_for_name,
        "RequestedForEmpID": requested_for_empid,
        "Req_JobTitle": req_job_title,
        "Req_JobCode": req_job_code,
        "Onshore_Offshore": onshore_offshore,
        "Hire Dates": expected_hire_date,
        "DirectReportsFlag": direct_reports,
        "PostingScope": posting_scope,
        "RequesterName": requester_name,
        "Hiring Need": hiring_need,
        "Business Justification": business_justification,
        "Job Family": job_family,
        "Level": level,
        "Business Unit": business_unit,
        "Quarter": quarter,
        "JD File": f"{next_req_id}.docx",
        "Date": datetime.now(),
        "Request ID": next_req_id,
        "WorkFlowLevel": 3
    }

    request_df = pd.read_excel(
        "Data.xlsx",
        sheet_name="Request"
    )

    request_df = pd.concat(
        [request_df, pd.DataFrame([request_row])],
        ignore_index=True
    )

    # Read WFDetails
    wf_details_df = pd.read_excel(
        "Data.xlsx",
        sheet_name="WFDetails"
    )

    wf_row = {
        "REQ": next_req_id,
        "WorkflowLevel": 3,
        "Date": datetime.now().strftime("%m/%d/%Y")
    }

    wf_details_df = pd.concat(
        [wf_details_df, pd.DataFrame([wf_row])],
        ignore_index=True
    )

    # Read Workflow sheet
    workflow_df = pd.read_excel(
        "Data.xlsx",
        sheet_name="Workflow"
    )

    # Save everything back
    with pd.ExcelWriter(
        "Data.xlsx",
        engine="openpyxl",
        mode="w"
    ) as writer:

        request_df.to_excel(
            writer,
            sheet_name="Request",
            index=False
        )

        workflow_df.to_excel(
            writer,
            sheet_name="Workflow",
            index=False
        )

        wf_details_df.to_excel(
            writer,
            sheet_name="WFDetails",
            index=False
        )

    # Save JD as Word document

    doc = Document()

    doc.add_heading(req_job_title, level=1)

    doc.add_heading("Job Description", level=2)

    # Replace generated_jd with your actual JD variable
    doc.add_paragraph(
        st.session_state.get("generated_jd", "")
    )

    doc_name = f"{next_req_id}.docx"

    # Create JD folder if missing
    os.makedirs("JD", exist_ok=True)

    doc_name = f"{next_req_id}.docx"
    doc_path = os.path.join("JD", doc_name)

    doc.save(doc_path)


def render_headcount_justification_tab():
    """
    Additional tab content for Submit Request for BU Review.
    Questions are based on the uploaded New Headcount Request Form.
    Existing submit/save/email logic is not changed by this function.
    """

    # Initialize defaults before widgets are created.
    # This avoids Streamlit widget/session_state conflicts.
    justification_defaults = {
        "budget_cost_reimbursable": "No",
        "budget_approved_in_workforce_plan": "No",
        "budget_status": "Under Budget",
        "budget_primary_objective": st.session_state.get("business_justification", ""),
        "budget_demand_utilization_change": st.session_state.get("hiring_need", ""),
        "budget_estimated_cost": "",
        "budget_contractor_conversion": "",
        "budget_additional_info": "",
        "location_offshore_possible": "Yes",
        "location_offshore_no_reason": "",
    }

    for key, default_value in justification_defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

    st.subheader("Justification and Budget")
    st.caption("Complete these questions before submitting the new headcount request.")

    st.radio(
        "Is this headcount for a cost reimbursable project (e.g., NAEP, billable to a contract)?",
        ["Yes", "No"],
        horizontal=True,
        key="budget_cost_reimbursable",
    )

    st.radio(
        "Was this headcount approved in the annual budgeting/workforce planning process with a hire date on or before the month you are submitting this request?",
        ["Yes", "No"],
        horizontal=True,
        key="budget_approved_in_workforce_plan",
    )

    st.radio(
        "In your last Financial Stewardship Meeting or monthly/quarterly budget meeting, was your Business Unit under budget or over budget?",
        ["Under Budget", "Over Budget"],
        horizontal=True,
        key="budget_status",
    )

    st.text_area(
        "What is the primary business objective of this position and benefit to ETS?",
        key="budget_primary_objective",
        max_chars=500,
        height=90,
        help="Examples: cost savings, increased revenue, new business opportunity, execution of a critical contract or strategy. 500-character limit.",
    )

    st.text_area(
        "What changes in demand or utilization justify the new headcount?",
        key="budget_demand_utilization_change",
        max_chars=300,
        height=80,
        help="300-character limit.",
    )

    st.text_area(
        "What is the estimated cost of this position for ETS (US$ or INR)?",
        key="budget_estimated_cost",
        max_chars=300,
        height=80,
        help="300-character limit.",
    )

    st.text_area(
        "If this request is for a Contractor/Temp conversion, provide the current contractor/temp bill rate and estimated salary as a full-time employee. (Optional)",
        key="budget_contractor_conversion",
        height=80,
    )

    st.text_area(
        "Additional information related to the position. (Optional)",
        key="budget_additional_info",
        height=80,
    )

    st.divider()
    st.subheader("Location Question")

    offshore_possible = st.radio(
        "Is this an offshore position? Can this position be offshored?",
        ["Yes", "No"],
        horizontal=True,
        key="location_offshore_possible",
    )

    if offshore_possible == "No":
        st.text_area(
            "If No selected, why?",
            key="location_offshore_no_reason",
            height=80,
        )

@st.dialog("Submit Request for BU Review", width="large")
def submit_request_dialog():

    show_process_flow(current_step=2)

    
    det_tab, justification_tab, jd_tab = st.tabs(
            ["📄 Details", "📋 Justification and Budget", "✏️ Edit JD"]
        )

    with justification_tab:
        render_headcount_justification_tab()

    with jd_tab:
        st.subheader("Final Job Description")

        final_jd = st.text_area(
            "Final JD",
            value=st.session_state.get("generated_jd", ""),
            height=350
        )

    with det_tab:
        req_df = pd.read_excel("Data.xlsx", sheet_name="Request")

        # Generate next Request ID
        last_req = (
            req_df["Request ID"]
            .astype(str)
            .str.replace("REQ", "", regex=False)
            .astype(int)
            .max()
        )

        next_req_id = f"REQ{last_req + 1:03d}"
        doc_name = f"{next_req_id}.docx"
        doc_path = os.path.join("JD", doc_name)

        col1, col2, col3 = st.columns(3)

        with col1:
            st.text_input("Request ID", value=next_req_id, disabled=True)
            st.text_input("Request Date", value=datetime.today().strftime("%m/%d/%Y"), disabled=True)
            st.text_input("Business Unit", value=st.session_state.get("business_unit", ""))
            st.text_input("Division", value="")
            requester_name = st.text_input("Requester Name", value="")
            requester_empid = st.text_input("Requester Emp ID", value="")
            requested_for_name = st.text_input("Requested For Name", value="")
            requested_for_empid = st.text_input("Requested For Emp ID", value="")

        with col2:

            req_job_title = st.text_input(
                "Req Job Title",
                value=st.session_state.get("job_family", "")
            )

            req_job_code = st.text_input(
                "Req Job Code",
                value=""
            )

            onshore_offshore = st.selectbox(
                "Onshore / Offshore",
                ["Onshore", "Offshore"]
            )

            expected_hire_date = st.date_input(
                "Expected Hire Date"
            )

            direct_reports = st.selectbox(
                "Direct Reports Flag",
                ["Yes", "No"]
            )
            posting_scope = st.text_input("Posting Scope")

            business_unit = st.session_state.get("business_unit", "") #st.text_input("Business Unit")

        with col3:
            hiring_need = st.text_area(
                "Hiring Need",
                value=st.session_state.get("hiring_need", "")
            )

            business_justification = st.text_area(
                "Business Justification",
                value=st.session_state.get("business_justification", "")
            )

            job_family = st.text_input(
                "Job Family",
                value=st.session_state.get("job_family", "")
            )

            level = st.text_input(
                "Level",
                value=st.session_state.get("job_level", "")
            )

            quarter = st.text_input(
                "Quarter",
                value=st.session_state.get("target_start_quarter", "")
            )


            # salary_min = st.text_input(
            #     "Min Salary",
            #     value= st.session_state.get("suggested_min_salary", "")
            # )
            # salary_max = st.text_input(
            #     "Max Salary",
            #     value=st.session_state.get("suggested_max_salary", "")
            # )

    # det_tab code ends

    if st.button(
        "Submit for BU Review",
        type="primary",
        width='stretch'
    ):

        # Existing logic remains unchanged:
        # 1. Generate Request ID
        # 2. Save JD to JD/REQxxx.docx
        # 3. Update Request sheet
        # 4. Update WFDetails sheet
        # 5. Send email

        save_request(
            next_req_id,
            requester_name,
            requester_empid,
            requested_for_name,
            requested_for_empid,
            req_job_title,
            req_job_code,
            onshore_offshore,
            expected_hire_date,
            direct_reports,
            posting_scope,
            hiring_need,
            business_justification,
            job_family,
            level,
            business_unit,
            quarter,
            doc_path)

        # send email
        # send_hiring_request_email(..
        
        # ✅ SEND EMAIL AFTER SAVE
        try:
            send_jd_email(
                docx_path=doc_path,
                recipient_email="shjoshi@ets.org",
                subject=f"JD Generated - {next_req_id}",
                body=f"""Hello,
        
        The JD for request ID {next_req_id} has been generated.
        
        Job Title: {req_job_title}
        Business Unit: {business_unit}
        
        Regards,
        JD Generator App"""
            )
        
            st.success(
                f"{next_req_id} submitted for BU Review and emailed successfully."
            )
        
        except Exception as e:
            st.warning(
                f"{next_req_id} submitted for BU Review, but email failed."
            )
            st.error(str(e))

        st.session_state.show_submit_dialog = False
        # st.success(f"{next_req_id} submitted for BU Review successfully.")
        
        # =====================================================
        # DOWNLOAD JD AFTER SUBMISSION
        # =====================================================
        submitted_jd_path = os.path.join("JD", f"{next_req_id}.docx")

        if os.path.exists(submitted_jd_path):
            with open(submitted_jd_path, "rb") as file:
                st.download_button(
                    label="Download JD",
                    data=file,
                    file_name=f"{next_req_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width="stretch"
                )
        else:
            st.warning("JD file was submitted, but the download file was not found.")

        
        # =====================================================
        # DOWNLOAD UPDATED DATA.XLSX AFTER SUBMISSION
        # =====================================================
        if os.path.exists("Data.xlsx"):
            with open("Data.xlsx", "rb") as excel_file:
                st.download_button(
                    label="Download Current Data.xlsx",
                    data=excel_file,
                    file_name="Data.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    width="stretch"
                )
        else:
            st.warning("Data.xlsx file was not found.")


        # st.rerun()
        # st.rerun()


def show():
    # =====================================================
    # CONFIG
    # =====================================================
    # col1, col2 = st.columns([1, 5])
    
    # with col1:
    #     st.image(r"C:\Users\JoshiShilpa\Documents\Shilpa\04_JD_Bot\ETS_Digital_RGB_Ink.jpeg", width=100)


    st.set_page_config(
        page_title="JD Generator",
        layout="wide"
    )

    # =====================================================
    # ACTIVE LEFT TAB HIGHLIGHTING
    # =====================================================
    st.markdown(
        """
        <style>
        /* Highlight selected Streamlit tab on the left panel */
        button[data-baseweb="tab"][aria-selected="true"] {
        *   background-color: #001F5B !impo*tant;
            color: white !im*ortant;
            border-radius:*8px 8px 0px 0px !important;
            font-weight: 700 !important;
        }

        /* Keep unselected tabs clean */
        button[data-baseweb="tab"][aria-selected="false"] {
       *    background-color: #F2F4F7 !imp*rtant;
            color: #001F5B *important;
*           border-radius: 8px 8px *px 0px !important;
            font-weight: 500 !important;
        }

        /* Optional: make tab text readable */
        button[data-baseweb="tab"] p {
            font-size: 15px *important;
        }
        </sty*e>
        """,
        unsafe_allow_html=True
    )

    show_process_flow()

    # Initialize once near the top of the app
    if "slt_sent" not in st.session_state:
        st.session_state.slt_sent = False

    if "generated_jd" not in st.session_state:
        st.session_state.generated_jd = None

    if "awaiting_review" not in st.session_state:
        st.session_state.awaiting_review = False

    if "submission_status" not in st.session_state:
        st.session_state.submission_status = None

    # if "salary_range" not in st.session_state:
    #     st.session_state.salary_range = None

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
        
    if "show_submit_dialog" not in st.session_state:
        st.session_state.show_submit_dialog = False

    if st.session_state.show_submit_dialog:
        submit_request_dialog()

    # =====================================================
    # AZURE OPENAI
    # =====================================================
    load_dotenv()
    client = AzureOpenAI(
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
    )

    deployment_name = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT")

    # =====================================================
    # HELPER FUNCTIONS
    # =====================================================

    def get_experience_range(level):
        mapping = {
            "I": "2-4 years experience",
            "II": "4-6 years experience",
            "III": "6-10 years senior-level experience",
            "IV": "10-15 years experience"
        }
        return mapping[level]


    def get_salary_range(job_family, level):
        """
        Replace with salary API/tool.
        For now returns sample values.
        """

        salary_map = {
            ("Software Engineer", "I"): (85000, 105000),
            ("Software Engineer", "II"): (105000, 130000),
            ("Software Engineer", "III"): (130000, 170000),
            ("Software Engineer", "IV"): (170000, 220000),

            ("Data Engineer", "I"): (90000, 115000),
            ("Data Engineer", "II"): (115000, 140000),
            ("Data Engineer", "III"): (140000, 180000),
            ("Data Engineer", "IV"): (180000, 230000),

            ("Product Manager", "I"): (95000, 120000),
            ("Product Manager", "II"): (120000, 145000),
            ("Product Manager", "III"): (145000, 185000),
            ("Product Manager", "IV"): (185000, 240000),

            ("Architect", "I"): (110000, 135000),
            ("Architect", "II"): (135000, 170000),
            ("Architect", "III"): (170000, 220000),
            ("Architect", "IV"): (220000, 280000)
        }

        return salary_map.get((job_family, level), (100000, 150000))


    def get_reference_jds(job_family, level):
        """
        Replace with web search tool.
        Ideally fetch 3 relevant JD urls.
        """

        experience = get_experience_range(level)

        return [
                "https://resumeworded.com/product-manager-resume-examples",
                "https://resume.io/resume-examples/architect",
                "https://enhancv.com/resume-examples/software-engineer/",
                "https://resumeworded.com/data-engineer-resume-examples"
        ], experience

    # =====================================================
    # MARKDOWN
    # =====================================================

    # st.markdown("""
    # <style>

    # /* Navy blue buttons */
    # .stButton > button[kind="primary"],
    # .stFormSubmitButton > button {
    #     background-color: #001F5B !important;
    #     color: white !important;
    # }

    # /* Slider handle */
    # .stSlider [role="slider"] {
    #     background-color: #001F5B !important;
    # }

    # /* Slider track */
    # .stSlider div[data-baseweb="slider"] div {
    #     color: #001F5B !important;
    # }

    # /* Active (selected) slider range */
    # .stSlider [data-baseweb="slider"] div[class*="track"] {
    #     background-color: #001F5B !important;
    # }

    # </style>
    # """, unsafe_allow_html=True)

    # =====================================================
    # LAYOUT
    # =====================================================

    left_col, right_col = st.columns([1, 2])

    # =====================================================
    # LEFT PANEL
    # =====================================================

    with left_col:


        tab1, tab2 = st.tabs([
            "📝 Create JD",
            "🤖 Chat with AI"
        ])

        with tab1:
            st.header("Hiring Request")

            what_need = st.text_area(
                "What do you need to hire",
                height=100
            )

            business_justification = st.text_area(
                "Business Justification",
                height=150
            )

            col1, col2 = st.columns(2)

            with col1:
                job_family = st.selectbox(
                    "Job Family",
                    [
                        "Software Engineer",
                        "Data Engineer",
                        "Architect",
                        "Product Manager"
                    ]
                )

            with col2:
                level = st.selectbox(
                    "Level",
                    ["I", "II", "III", "IV"]
                )

            col3, col4 = st.columns(2)

            with col3:
                business_unit = st.selectbox(
                    "Business Unit",
                    [
                        "Technology Delivery",
                        "Enterprise Products",
                        "Data Operation"
                    ]
                )

            with col4:
                target_start_quarter = st.selectbox(
                    "Target Start Quarter",
                    [
                        "2026-Q3",
                        "2026-Q4",
                        "2027-Q1"
                    ]
                )

            generate = st.button(
                "Generate JD",
                type="primary",
                width='stretch'
            ) # end of tab 1 in the left panel

        with tab2:

            # Show conversation only
            for message in st.session_state.chat_history:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            prompt = st.chat_input(
                "Describe the role you want to create..."
            )

            if prompt:

                st.session_state.chat_history.append({
                    "role": "user",
                    "content": prompt
                })

                response = client.chat.completions.create(
                    model= deployment_name,
                    messages=[
                        {
                            "role": "system",
                            "content": """
                            You are an HR recruiter.
                            Generate and refine Job Descriptions.
                            """
                        }
                    ] + st.session_state.chat_history
                )

                jd_response = response.choices[0].message.content

                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": "JD updated. Please review the preview panel."
                })

                # Store JD separately
                st.session_state.generated_jd = jd_response

                st.rerun()
                
    # =====================================================
    # RIGHT PANEL
    # =====================================================

    with right_col:

        # ----------------------------
        # Session state initialization
        # ----------------------------
        if "generated_jd" not in st.session_state:
            st.session_state.generated_jd = None

        if "awaiting_review" not in st.session_state:
            st.session_state.awaiting_review = False

        if "submission_status" not in st.session_state:
            st.session_state.submission_status = None

        # if "salary_range" not in st.session_state:
        #     st.session_state.salary_range = None

        # ----------------------------
        # Generate JD
        # ----------------------------
        if generate:

            suggested_min_salary, suggested_max_salary = 0,0

            reference_sites, experience_range = get_reference_jds(
                job_family,
                level
            )

            st.session_state.hiring_need = what_need
            st.session_state.business_unit = business_unit
            st.session_state.business_justification = business_justification
            st.session_state.job_family = job_family
            st.session_state.job_level = level
            st.session_state.business_unit = business_unit
            st.session_state.target_start_quarter = target_start_quarter

            prompt = f"""
            You are an HR Admin.

            Create a professional Job Description.

            Include sections:

            1. Role Summary
            2. Key Responsibilities
            3. Qualifications
            4. Context Summary

            Hiring Need:
            {what_need}

            Business Justification:
            {business_justification}

            Job Family:
            {job_family}

            Level:
            {level}

            Expected Experience:
            {experience_range}

            Business Unit:
            {business_unit}

            Target Start Quarter:
            {target_start_quarter}

            

            Reference Job Description Sources:

            1. {reference_sites[0]}
            2. {reference_sites[1]}
            3. {reference_sites[2]}

            Use the referenced job descriptions only as inspiration.
            Generate a unique JD.

            Do not include suggestions.
            """

            with st.spinner("Generating JD..."):
                response = client.chat.completions.create(
                model=deployment_name,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an HR Admin who writes professional job descriptions."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.7,
                max_completion_tokens=2000
            )

            jd = response.choices[0].message.content

            st.session_state.generated_jd = jd
            # st.session_state.suggested_min_salary = suggested_min_salary
            # st.session_state.suggested_max_salary = suggested_max_salary

            # st.session_state.salary_range = (
            #     suggested_min_salary,
            #     suggested_max_salary
            # )

            st.session_state.submission_status = None
            st.session_state.awaiting_review = False

        # ----------------------------
        # Show JD Form
        # ----------------------------
        if st.session_state.generated_jd:

            jd = st.session_state.generated_jd
            # suggested_min_salary = st.session_state.suggested_min_salary
            # suggested_max_salary = st.session_state.suggested_max_salary

            st.header("Generated Job Description")

            with st.form("job_form"):

                tab1, tab2 = st.tabs(["Preview", "Edit"])

                with tab1:
                    with st.container(height=350):
                        st.markdown(jd)

                with tab2:
                    edited_jd = st.text_area(
                        "Edit Job Description",
                        value=jd,
                        height=350,
                        key="edit_jd"
                    )

                # st.subheader("Salary Range")

                # min_salary, max_salary = st.slider(
                #     "Select Salary Range",
                #     min_value=max(0, suggested_min_salary - 15000),
                #     max_value=suggested_max_salary + 15000,
                #     value=st.session_state.salary_range,
                #     step=1000
                # )

                # st.info(
                #     f"Recommended Range: "
                #     f"${suggested_min_salary:,} - ${suggested_max_salary:,}"
                # )

                col1, col2 = st.columns(2)
                with col1:
                    submitted = st.form_submit_button(
                        "Use this JD",
                        width='stretch'
                    )


            # Persist slider value
            # st.session_state.salary_range = (
            #     min_salary,
            #     max_salary
            # )

            # Ask review question after submit
            if submitted:
                st.session_state.awaiting_review = True
                st.session_state.submission_status = None

        # ----------------------------
        # Review Confirmation
        # ----------------------------
        if st.session_state.awaiting_review:

            # st.divider()

            st.warning(
                "Before submitting, please confirm that the Job Description has been reviewed."
            )

            review_choice = st.radio(
                "Has the JD been reviewed?",
                ["Yes", "No"],
                index  = 1,
                horizontal=True,
                key="jd_review_confirmation"
            )

            if st.button(
                "Confirm JD",
                type="primary",
                width='stretch'
            ):

                if review_choice == "Yes":
                    st.session_state.show_submit_dialog = True
                    with st.spinner("Submitting for review..."):
                        submit_request_dialog()

                else:

                    st.session_state.submission_status = "not_reviewed"
                st.session_state.awaiting_review = False
                st.rerun()


            docx_file = create_docx(jd)
            st.download_button(
                label="Download JD",
                data=docx_file,
                file_name="Job_Description.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )


                
                

        elif st.session_state.submission_status == "not_reviewed":
            st.warning("⚠️ Review the JD before submitting for review")


